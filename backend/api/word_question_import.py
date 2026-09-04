"""
Parse and generate Word (.docx) question files for homework / lessons / banks.

Format (copy this file and stack more blocks under each other):

    سؤال
    نص السؤال

    أ) الخيار الأول
    ب) الخيار الثاني
    ج) الخيار الثالث
    د) الخيار الرابع

    الإجابة الصحيحة: ب
    الشرح: اختياري

    قطعة
    نص القطعة المشتركة

    سؤال
    أول سؤال تحت القطعة
    أ) ...
    ب) ...
    ج) ...
    د) ...
    الإجابة الصحيحة: أ
    الشرح: اختياري

    سؤال
    ثاني سؤال تحت نفس القطعة
    ...

    نهاية القطعة
"""

from __future__ import annotations

import html
import io
import re

OPTION_IDS = ("a", "b", "c", "d")
OPTION_AR = {"a": "أ", "b": "ب", "c": "ج", "d": "د"}

_LETTER_TO_ID = {
    "أ": "a",
    "ا": "a",
    "إ": "a",
    "آ": "a",
    "a": "a",
    "A": "a",
    "ب": "b",
    "b": "b",
    "B": "b",
    "ج": "c",
    "c": "c",
    "C": "c",
    "د": "d",
    "d": "d",
    "D": "d",
}

_RE_QUESTION_HEAD = re.compile(
    r"^(?:[#=\-*_ ]*)(سؤال|question)(?:\s*[:：-]?\s*\d+)?(?:[#=\-*_ ]*)?$",
    re.IGNORECASE,
)
_RE_PASSAGE_HEAD = re.compile(
    r"^(?:[#=\-*_ ]*)(قطعة|passage)(?:\s*[:：-]?\s*\d+)?(?:[#=\-*_ ]*)?$",
    re.IGNORECASE,
)
_RE_END_PASSAGE = re.compile(
    r"^(?:[#=\-*_ ]*)(نهاية\s*القطعة|end\s*passage)(?:[#=\-*_ ]*)?$",
    re.IGNORECASE,
)
_RE_OPTION = re.compile(
    r"^([أاإآبجدAaBbCcDd])\s*[)）\]\-.\-:：]\s*(.*)$",
)
_RE_CORRECT = re.compile(
    r"^(?:الإجابة\s*الصحيحة|الإجابة|الصحيح|correct(?:\s*answer)?)\s*[:：]\s*(.*)$",
    re.IGNORECASE,
)
_RE_EXPLAIN = re.compile(
    r"^(?:الشرح|التوضيح|التفسير|explanation)\s*[:：]\s*(.*)$",
    re.IGNORECASE,
)
_RE_SEPARATOR = re.compile(r"^[-_=*#]{3,}$")
_RE_START_MARKER = re.compile(r"ابدأ\s*الأسئلة", re.IGNORECASE)


def _norm(text: str) -> str:
    if not text:
        return ""
    t = text.replace("\u200f", "").replace("\u200e", "").replace("\xa0", " ")
    t = t.replace("ى", "ي")
    return re.sub(r"\s+", " ", t).strip()


def _letter_id(raw: str) -> str | None:
    raw = _norm(raw)
    if not raw:
        return None
    ch = raw[0]
    return _LETTER_TO_ID.get(ch)


def _to_html(parts: list[str]) -> str:
    chunks = [_norm(p) for p in parts if _norm(p)]
    if not chunks:
        return ""
    if len(chunks) == 1:
        return html.escape(chunks[0])
    return "".join(f"<p>{html.escape(p)}</p>" for p in chunks)


def _empty_question() -> dict:
    return {
        "question_parts": [],
        "answers": {oid: [] for oid in OPTION_IDS},
        "correct": None,
        "explanation_parts": [],
        "current_option": None,
        "in_explanation": False,
    }


def _empty_passage() -> dict:
    return {
        "type": "passage",
        "passage_parts": [],
        "questions": [],
        "current": None,
        "phase": "body",  # body | question | closed
        "in_explanation": False,
    }


def _empty_single() -> dict:
    q = _empty_question()
    q["type"] = "single"
    return q


def _finalize_question(q: dict, index_label: str) -> tuple[dict, list[str]]:
    errors: list[str] = []
    text = _to_html(q.get("question_parts") or [])
    if not text:
        errors.append(f"{index_label}: نص السؤال فارغ.")

    answers = []
    missing = []
    for oid in OPTION_IDS:
        ans_text = _to_html(q.get("answers", {}).get(oid) or [])
        if not ans_text:
            missing.append(OPTION_AR[oid])
        answers.append(
            {
                "answer_id": oid,
                "text": ans_text,
                "is_correct": q.get("correct") == oid,
            }
        )
    if missing:
        errors.append(
            f"{index_label}: ناقص الاختيار ({'، '.join(missing)}). يجب وجود أ ب ج د."
        )

    correct = q.get("correct")
    if not correct:
        errors.append(f"{index_label}: لم تُحدد الإجابة الصحيحة.")
    elif not any(a["is_correct"] for a in answers):
        errors.append(f"{index_label}: الإجابة الصحيحة لا تطابق أي اختيار.")

    item = {
        "question": text,
        "explanation": _to_html(q.get("explanation_parts") or []) or None,
        "answers": answers,
    }
    return item, errors


def _append_text(target_list: list[str], text: str) -> None:
    t = text.strip()
    if t:
        target_list.append(t)


def _apply_line_to_question(q: dict, kind: str, payload) -> None:
    if kind == "option":
        oid, rest = payload
        q["current_option"] = oid
        q["in_explanation"] = False
        _append_text(q["answers"][oid], rest)
        return
    if kind == "correct":
        q["in_explanation"] = False
        q["current_option"] = None
        q["correct"] = payload
        return
    if kind == "explain":
        q["in_explanation"] = True
        q["current_option"] = None
        _append_text(q["explanation_parts"], payload)
        return
    # text
    if q.get("in_explanation"):
        _append_text(q["explanation_parts"], payload)
    elif q.get("current_option"):
        _append_text(q["answers"][q["current_option"]], payload)
    else:
        _append_text(q["question_parts"], payload)


def _classify(line: str, style: str) -> tuple[str, object]:
    raw = _norm(line)
    if not raw:
        return "empty", None
    if _RE_SEPARATOR.match(raw):
        return "separator", None
    if _RE_END_PASSAGE.match(raw):
        return "end_passage", None
    if _RE_PASSAGE_HEAD.match(raw) or (
        style.lower().startswith("heading 1") and _RE_PASSAGE_HEAD.match(raw)
    ):
        return "passage_head", None
    if _RE_QUESTION_HEAD.match(raw):
        heading2 = style.lower().startswith("heading 2")
        return ("sub_question_head" if heading2 else "question_head"), None
    m_opt = _RE_OPTION.match(raw)
    if m_opt:
        oid = _letter_id(m_opt.group(1))
        if oid:
            return "option", (oid, m_opt.group(2) or "")
    m_c = _RE_CORRECT.match(raw)
    if m_c:
        return "correct", _letter_id(m_c.group(1) or "")
    m_e = _RE_EXPLAIN.match(raw)
    if m_e:
        return "explain", m_e.group(1) or ""
    if _RE_START_MARKER.search(raw):
        return "start_marker", None
    return "text", raw


def _iter_docx_lines(file_obj) -> list[tuple[str, str]]:
    from docx import Document

    if hasattr(file_obj, "seek"):
        try:
            file_obj.seek(0)
        except Exception:
            pass
    doc = Document(file_obj)
    rows: list[tuple[str, str]] = []
    for p in doc.paragraphs:
        style = ""
        try:
            style = p.style.name if p.style is not None else ""
        except Exception:
            style = ""
        rows.append((p.text or "", style or ""))
    # Tables: each cell paragraph, in case someone pastes a table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    style = ""
                    try:
                        style = p.style.name if p.style is not None else ""
                    except Exception:
                        style = ""
                    rows.append((p.text or "", style or ""))
    return rows


def parse_docx_file(file_obj) -> dict:
    """
    Return {
      items: [...],
      errors: [...],
      warnings: [...],
      summary: {single, passage, questions}
    }
    """
    try:
        lines = _iter_docx_lines(file_obj)
    except Exception as exc:
        return {
            "items": [],
            "errors": [f"تعذر قراءة ملف وورد. تأكد أنه بصيغة .docx ({exc})"],
            "warnings": [],
            "summary": {"single": 0, "passage": 0, "questions": 0},
        }

    blocks: list[dict] = []
    current = None
    started = False
    warnings: list[str] = []

    def flush_passage_question():
        if current and current.get("type") == "passage" and current.get("current"):
            current["questions"].append(current["current"])
            current["current"] = None

    def flush_current():
        nonlocal current
        if current is None:
            return
        if current.get("type") == "passage":
            flush_passage_question()
        blocks.append(current)
        current = None

    for text, style in lines:
        kind, payload = _classify(text, style)
        if kind in ("empty", "separator"):
            continue
        if kind == "start_marker":
            started = True
            continue

        if kind in ("question_head", "sub_question_head", "passage_head"):
            started = True

        if not started:
            continue

        if kind == "passage_head":
            flush_current()
            current = _empty_passage()
            continue

        if kind == "end_passage":
            if current and current.get("type") == "passage":
                flush_passage_question()
                current["phase"] = "closed"
            continue

        if kind in ("question_head", "sub_question_head"):
            in_open_passage = (
                current
                and current.get("type") == "passage"
                and current.get("phase") != "closed"
            )
            if in_open_passage:
                flush_passage_question()
                current["phase"] = "question"
                current["current"] = _empty_question()
            else:
                flush_current()
                current = _empty_single()
            continue

        if current is None:
            if kind == "text":
                warnings.append("تم تجاهل نص قبل أول عنوان «سؤال» أو «قطعة».")
            continue

        if current.get("type") == "single":
            if kind == "correct" and not payload:
                warnings.append("سطر الإجابة الصحيحة بلا حرف (أ/ب/ج/د).")
            _apply_line_to_question(current, kind, payload)
            continue

        # passage
        if current.get("phase") == "closed":
            warnings.append("نص بعد «نهاية القطعة» تم تجاهله حتى سؤال/قطعة جديد.")
            continue

        if current.get("phase") == "question" and current.get("current"):
            if kind == "correct" and not payload:
                warnings.append("سطر الإجابة الصحيحة بلا حرف (أ/ب/ج/د).")
            _apply_line_to_question(current["current"], kind, payload)
        else:
            if kind == "text":
                _append_text(current["passage_parts"], payload)
            elif kind in ("option", "correct", "explain"):
                warnings.append(
                    "وُجدت اختيارات أو إجابة قبل سؤال القطعة؛ أضف عنوان «سؤال» تحت القطعة أولاً."
                )

    flush_current()

    items = []
    errors: list[str] = []
    single_count = 0
    passage_count = 0
    question_count = 0

    for i, block in enumerate(blocks, start=1):
        if block.get("type") == "single":
            item, errs = _finalize_question(block, f"سؤال {i}")
            errors.extend(errs)
            items.append({"type": "single", **item})
            single_count += 1
            question_count += 1
            continue

        passage_html = _to_html(block.get("passage_parts") or [])
        if not passage_html:
            errors.append(f"قطعة {i}: نص القطعة فارغ.")
        sub_items = []
        subs = block.get("questions") or []
        if not subs:
            errors.append(f"قطعة {i}: أضف سؤالاً واحداً على الأقل تحت القطعة.")
        for j, sub in enumerate(subs, start=1):
            sub_item, sub_errs = _finalize_question(sub, f"قطعة {i} / سؤال {j}")
            errors.extend(sub_errs)
            sub_items.append(sub_item)
            question_count += 1
        items.append(
            {
                "type": "passage",
                "passage_text": passage_html,
                "questions": sub_items,
            }
        )
        passage_count += 1

    if not items:
        errors.append(
            "لم يُعثر على أسئلة. استخدم عنوان «سؤال» أو «قطعة» كما في ملف المثال."
        )

    return {
        "items": items,
        "errors": errors,
        "warnings": warnings,
        "summary": {
            "single": single_count,
            "passage": passage_count,
            "questions": question_count,
        },
    }


def parsed_items_to_api_payloads(items: list[dict], lesson_id: str) -> list[dict]:
    payloads = []
    for item in items:
        if item.get("type") == "passage":
            payloads.append(
                {
                    "lesson": lesson_id,
                    "question_type": "passage",
                    "question": "(قطعة)",
                    "answers": [],
                    "passage_text": item.get("passage_text") or "",
                    "passage_questions": [
                        {
                            "question": q.get("question") or "",
                            "explanation": q.get("explanation") or None,
                            "answers": q.get("answers") or [],
                        }
                        for q in (item.get("questions") or [])
                    ],
                }
            )
        else:
            payloads.append(
                {
                    "lesson": lesson_id,
                    "question_type": "single",
                    "question": item.get("question") or "",
                    "explanation": item.get("explanation") or None,
                    "answers": item.get("answers") or [],
                    "passage_text": "",
                    "passage_questions": [],
                }
            )
    return payloads


def _set_rtl_paragraph(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)


def _set_run_rtl(run) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), "Traditional Arabic")
    rFonts.set(qn("w:ascii"), "Traditional Arabic")
    rFonts.set(qn("w:hAnsi"), "Traditional Arabic")


def _add_p(doc, text, *, style=None, bold=False, size=None, color=None, space_after=80):
    from docx.shared import Pt, RGBColor, Twips

    p = doc.add_paragraph()
    if style:
        try:
            p.style = style
        except Exception:
            pass
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    _set_rtl_paragraph(p)
    _set_run_rtl(run)
    p.paragraph_format.space_after = Twips(space_after)
    return p


def build_template_document():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Cm

    doc = Document()
    section = doc.sections[0]
    section.right_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)

    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    sectPr.append(bidi)

    styles = doc.styles
    try:
        styles["Normal"].font.name = "Traditional Arabic"
        styles["Normal"].font.size = Pt(14)
    except Exception:
        pass

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("ملف المثال — رفع أسئلة الواجب / الدرس / البنك")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = RGBColor(0, 95, 95)
    _set_rtl_paragraph(title)
    _set_run_rtl(tr)

    intro_lines = [
        "انسخ هذا الملف كما هو، ثم احذف الأمثلة بعد فهم الشكل، وأضف أسئلتك أسفل بعضها بنفس العناوين.",
        "القواعد:",
        "١) ابدأ كل سؤال عادي بسطر عنوان وحده: سؤال",
        "٢) الاختيارات الأربعة بهذا الشكل بالضبط: أ)   ب)   ج)   د)",
        "٣) سطر الإجابة الصحيحة حرف واحد فقط، مثال: الإجابة الصحيحة: ب",
        "٤) سطر الشرح اختياري. إن لم يوجد شرح احذف السطر.",
        "٥) القطعة: عنوان قطعة ثم نصها، ثم عنوان سؤال لكل سؤال تحتها، ثم نهاية القطعة",
        "٦) لا ترفع صور داخل الملف. أضف الصور لاحقاً من صفحة إدارة الأسئلة إن لزم.",
        "٧) ابدأ كتابة الأسئلة بعد السطر التالي.",
        "٨) أضف الأسئلة أسفل بعضها بنفس الشكل، واحذف الأمثلة قبل الرفع إن أردت ملفاً فيه أسئلتك فقط.",
    ]
    for line in intro_lines:
        _add_p(doc, line, size=13, color=(60, 60, 60), space_after=60)

    start = doc.add_paragraph()
    sr = start.add_run("—— ابدأ الأسئلة ——")
    sr.bold = True
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0, 128, 128)
    _set_rtl_paragraph(start)
    _set_run_rtl(sr)

    def add_heading_ar(text, heading="Heading 1"):
        p = doc.add_paragraph()
        try:
            p.style = heading
        except Exception:
            pass
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18 if heading == "Heading 1" else 16)
        run.font.color.rgb = RGBColor(0, 95, 95)
        _set_rtl_paragraph(p)
        _set_run_rtl(run)
        return p

    def add_options(opts, correct, explanation):
        for oid in OPTION_IDS:
            _add_p(doc, f"{OPTION_AR[oid]}) {opts[oid]}", size=14, space_after=40)
        _add_p(
            doc,
            f"الإجابة الصحيحة: {OPTION_AR[correct]}",
            bold=True,
            size=14,
            color=(0, 100, 70),
            space_after=40,
        )
        if explanation:
            _add_p(doc, f"الشرح: {explanation}", size=13, color=(80, 80, 80), space_after=120)

    # Example 1 — single
    add_heading_ar("سؤال")
    _add_p(doc, "ما ناتج العملية: ٥ + ٧ ؟", size=15, space_after=80)
    add_options(
        {
            "a": "١٠",
            "b": "١١",
            "c": "١٢",
            "d": "١٣",
        },
        "c",
        "٥ + ٧ = ١٢",
    )

    # Example 2 — single without explanation
    add_heading_ar("سؤال")
    _add_p(
        doc,
        "اختر الكلمة الأقرب معنىً إلى: شجاع",
        size=15,
        space_after=80,
    )
    add_options(
        {
            "a": "جبان",
            "b": "مقدام",
            "c": "كسول",
            "d": "حزين",
        },
        "b",
        None,
    )

    # Example 3 — passage with two questions
    add_heading_ar("قطعة")
    _add_p(
        doc,
        "قرأ أحمد قصة قصيرة عن طالب اجتهد طوال العام، فحصد نتيجة مرتفعة في اختبار القدرات، ثم ساعد زملاءه على المراجعة قبل الاختبار المحوسب.",
        size=15,
        space_after=80,
    )

    add_heading_ar("سؤال", heading="Heading 2")
    _add_p(doc, "ما الصفة الأبرز لأحمد في القطعة؟", size=15, space_after=80)
    add_options(
        {
            "a": "الكسل",
            "b": "الاجتهاد",
            "c": "الغضب",
            "d": "النسيان",
        },
        "b",
        "النص يذكر أنه اجتهد طوال العام.",
    )

    add_heading_ar("سؤال", heading="Heading 2")
    _add_p(doc, "ماذا فعل أحمد بعد حصوله على نتيجة مرتفعة؟", size=15, space_after=80)
    add_options(
        {
            "a": "ترك الدراسة",
            "b": "سافر فوراً",
            "c": "ساعد زملاءه على المراجعة",
            "d": "غيّر تخصصه",
        },
        "c",
        "القطعة تقول إنه ساعد زملاءه على المراجعة قبل الاختبار المحوسب.",
    )

    endp = doc.add_paragraph()
    try:
        endp.style = "Heading 1"
    except Exception:
        pass
    er = endp.add_run("نهاية القطعة")
    er.bold = True
    er.font.size = Pt(16)
    er.font.color.rgb = RGBColor(120, 40, 40)
    _set_rtl_paragraph(endp)
    _set_run_rtl(er)

    return doc


def build_template_bytes() -> bytes:
    doc = build_template_document()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def write_template_file(path: str) -> str:
    doc = build_template_document()
    doc.save(path)
    return path
